from docx import Document
from docx.shared import Inches

document = Document()

# Profile picture
document.add_picture(
    'me.jpg',
    width=Inches(2.0)
)

# Name, Phone, Email details

name = input('What is your name?')
phone_number = input('What is your phone number?')
email = input('What is your email?')

# name = 'Olaki Edwin'
# phone_number = '0787499401'
# email = 'olakiedwin@gmail.com'

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself?')
)

document.save('cv.docx')
